# -*- coding: utf-8 -*-
import argparse
import io
import mimetypes
import os
import email
import re
import pyfiglet
import asyncio
import aiofiles
from email.header import decode_header
import docx
import pandas as pd
import xml.etree.ElementTree as ET
import time
import fitz
import subprocess
from pandas.io.sas.sas_constants import magic


class EmailProcessor:
    def __init__(self, file_path, log_file, error_file, keywords, output_folder):
        self.file_path = file_path
        self.log_file = log_file
        self.error_file = error_file
        self.keywords = keywords
        self.output_folder = output_folder

    def sanitize_filename(self, filename):
        valid_filename = re.sub(r'[\/:*?"<>|]', '_', filename)
        return valid_filename

    def decode_subject(self, subject):
        decoded = decode_header(subject)
        return decoded[0][0].decode(decoded[0][1]) if decoded[0][1] else decoded[0][0]

    def get_file_extension(self, filename, content):
        _, extension = os.path.splitext(filename)
        if extension:
            return extension.lstrip('.').lower()

        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            extension = mimetypes.guess_extension(mime_type)
            if extension:
                return extension.lstrip('.').lower()

        try:
            if content:
                mime_type = magic.Magic(mime=True).from_buffer(content)
                extension = mimetypes.guess_extension(mime_type)
                return extension.lstrip('.').lower() if extension else None
        except Exception:
            return None

    def decode_content(self, part):
        content_type = part.get_content_type()
        if "text/plain" in content_type:
            try:
                charset = part.get_content_charset()
                content = part.get_payload(decode=True)
                if charset:
                    content = content.decode(charset)
                return content
            except TypeError:
                return "Failed to decode content"
        elif "text/html" in content_type:
            return "HTML content:\n" + part.get_payload()
        else:
            return "Non-text content. Content type: " + content_type

    async def process_part(self, part, folder_path):
        try:
            content = self.decode_content(part)
            found_keywords = []
            for keyword in self.keywords:
                if keyword.lower() in content.lower():
                    found_keywords.append(keyword)
            if found_keywords:
                filename = self.file_path
                for keyword in found_keywords:
                    await self.log_found_keyword(filename, keyword)
        except Exception as e:
            await self.log_error(f"Помилка з файлом: {self.file_path}\nПомилка: {str(e)}")

    async def save_attachments(self, part, folder_path):
        try:
            if part.get_content_maintype() == "multipart":
                return
            filename = part.get_filename()
            if not filename:
                return

            decoded_filename, charset = decode_header(filename)[0]
            if charset:
                decoded_filename = decoded_filename.decode(charset)

            extension = self.get_file_extension(decoded_filename, part.get_payload(decode=True))
            if extension:
                sanitized_filename = self.sanitize_filename(decoded_filename)
                attachments_dir = os.path.join(folder_path, "attachments", extension)
                os.makedirs(attachments_dir, exist_ok=True)

                payload = part.get_payload(decode=True)

                content = None
                if extension == "txt":
                    content = payload.decode("utf-8")
                elif extension == "docx":
                    doc = docx.Document(io.BytesIO(payload))
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                elif extension == "pdf":
                    pdf_document = fitz.open(stream=io.BytesIO(payload))
                    content = ""
                    for page_num in range(len(pdf_document)):
                        page = pdf_document.load_page(page_num)
                        content += page.get_text()
                # elif extension == "xlsx":
                #     df = pd.read_excel(io.BytesIO(payload))
                #     content = df.to_string(index=False)
                elif extension == "xml":
                    tree = ET.ElementTree(ET.fromstring(payload))
                    root = tree.getroot()
                    content = ET.tostring(root, encoding="utf-8").decode("utf-8")
                elif extension == "csv":
                    df = pd.read_csv(io.BytesIO(payload))
                    content = df.to_string(index=False)

                elif extension in ["js", "css", "html", "json", "tsv"]:
                    content = payload.decode("utf-8")

                if content:
                    found_keywords = []
                    for keyword in self.keywords:
                        if keyword.lower() in content.lower():
                            found_keywords.append(keyword)

                    if found_keywords:
                        filepath = os.path.join(attachments_dir, sanitized_filename)
                        async with aiofiles.open(filepath, "wb") as attachment_file:
                            await attachment_file.write(payload)

                        for keyword in found_keywords:
                            await self.log_found_keyword(filepath, keyword)

        except Exception as e:
            await self.log_error(f"Помилка при збереженні вкладення: {str(e)}\n")
    async def log_found_keyword(self, filename, keyword):
        log_file_path = self.log_file
        drive, path = os.path.splitdrive(filename)
        dirs, filename = os.path.split(path)
        _, dirs = os.path.split(dirs)
        new_filename = os.path.join(os.path.sep, dirs, filename)

        async with aiofiles.open(log_file_path, "a", encoding="utf-8") as log:
            await log.write(f"Файл: {new_filename}, Ключевое слово: {keyword}\n")

    async def log_error(self, error_message):
        error_file_path = self.error_file
        async with aiofiles.open(error_file_path, "a", encoding="utf-8") as err_log:
            await err_log.write(error_message)

    async def process_email(self, folder_path, save_attachments=False):
        try:
            async with aiofiles.open(self.file_path, "rb") as file:
                data = await file.read()
                msg = email.message_from_bytes(data)

            tasks = []
            for part in msg.walk():
                if part.get_content_maintype() == "multipart":
                    continue
                task = self.process_part(part, folder_path)
                tasks.append(task)

                if save_attachments:
                    attachment_task = self.save_attachments(part, folder_path)
                    tasks.append(attachment_task)

            await asyncio.gather(*tasks)

        except Exception as e:
            await self.log_error(f"Помилка обробки електронної пошти в файлі: {self.file_path}\n")
            await self.log_error(f"Помилка: {str(e)}\n")

    async def process_file(self, file_path):
        try:
            extension = self.get_file_extension(file_path, None)
            content = None

            if extension == "txt":
                async with aiofiles.open(file_path, "r", encoding="utf-8") as text_file:
                    content = await text_file.read()
            elif extension == "docx":
                doc = docx.Document(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            elif extension == "pdf":
                pdf_document = fitz.open(file_path)
                num_pages = pdf_document.page_count
                content = ""

                for page_num in range(num_pages):
                    page = pdf_document.load_page(page_num)
                    content += page.get_text()

            elif extension == "csv":
                df = pd.read_csv(file_path)
                content = df.to_string(index=False)
            elif extension == "xlsx":
                df = pd.read_excel(file_path)
                content = df.to_string(index=False)

            elif extension == "xml":
                tree = ET.parse(file_path)
                root = tree.getroot()
                content = ET.tostring(root, encoding="utf-8").decode("utf-8")

            if content:
                found_keywords = []
                for keyword in self.keywords:
                    if keyword.lower() in content.lower():
                        found_keywords.append(keyword)

                if found_keywords:
                    for keyword in found_keywords:
                        await self.log_found_keyword(file_path, keyword)

        except Exception as e:
            await self.log_error(f"Помилка з файлом: {file_path}\nПомилка: {str(e)}")


async def search_keywords_in_emails(folder_path, log_file, error_file, keywords, output_folder, save_attachments=False):
    try:
        processed_files = set()
        os.makedirs(output_folder, exist_ok=True)
        tasks = []
        for root, dirs, files in os.walk(folder_path):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                if file_path not in processed_files:
                    if file_name.endswith(".eml"):
                        email_processor = EmailProcessor(file_path, log_file, error_file, keywords, output_folder)
                        task = asyncio.create_task(email_processor.process_email(output_folder, save_attachments))
                        tasks.append(task)
                    else:
                        email_processor = EmailProcessor(file_path, log_file, error_file, keywords, output_folder)
                        task = asyncio.create_task(email_processor.process_file(file_path))
                        tasks.append(task)
                    processed_files.add(file_path)

            await asyncio.gather(*tasks)

        print(f"Оброблено файлів: {len(processed_files)} в папці: {folder_path}")
    except Exception as e:
        print(f"Помилка в search_keywords_in_emails: {str(e)}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="пошук приколів")
    parser.add_argument("-f", "--folder", type=str, help="Шлях до листів")
    parser.add_argument("-l", "--log", type=str, help="Шлях для збереження лог-файлу")
    parser.add_argument("-e", "--error", type=str, help="Шлях для збереження лог-файлу з помилками")
    parser.add_argument("-k", "--keywords-file", type=str, help="Шлях до словника")
    parser.add_argument("-o", "--output-folder", type=str, help="Шлях до теки для збереження вкладень")
    parser.add_argument("-a", "--attachments", action='store_true',
                        help="Завантажити вкладення (за замовчуванням: False)")
    parser.add_argument("-d", "--extract", action='store_true', help="Запустити extract.py")

    args = parser.parse_args()
    folder_path = args.folder if args.folder else r"./"
    log_file = args.log if args.log else "log.txt"
    error_file = args.error if args.error else "errors.txt"
    output_folder = args.output_folder if args.output_folder else "attachments"
    save_attachments = args.attachments

    if args.extract and args.folder:
        extract_command = f"python extract.py {args.folder}"
        subprocess.Popen(extract_command, shell=True).wait()

    if args.keywords_file:
        with open(args.keywords_file, "r", encoding="utf-8") as keywords_file:
            keywords = [keyword.strip() for keyword in keywords_file.read().split(",")]

    start_time = time.time()
    os.makedirs(output_folder, exist_ok=True)
    text_to_display = "Email Parser"
    ascii_art = pyfiglet.figlet_format(text_to_display)
    print(ascii_art)

    asyncio.run(search_keywords_in_emails(folder_path, log_file, error_file, keywords, output_folder, save_attachments))

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(f"Час роботи скрипта: {elapsed_time} секунд.")
